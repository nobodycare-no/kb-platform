"""解析器单元测试（FR-B01）。"""
import io

import pytest

from app.services.parsers import ParserError, extract_text


def test_txt_utf8():
    text, ftype = extract_text("制度内容：差旅住宿标准。".encode("utf-8"), "a.txt")
    assert ftype == "txt" and "差旅" in text


def test_txt_gb18030_fallback():
    text, _ = extract_text("考勤制度：迟到扣款。".encode("gb18030"), "b.txt")
    assert "考勤" in text


def test_md_passthrough():
    content = "# 标题\n\n- 条目一\n- 条目二\n"
    text, ftype = extract_text(content.encode("utf-8"), "c.md")
    assert ftype == "md" and "# 标题" in text and "条目二" in text


def test_unsupported_extension_rejected():
    with pytest.raises(ParserError, match="不支持的文件类型"):
        extract_text(b"data", "virus.exe")


def test_empty_file_rejected():
    with pytest.raises(ParserError, match="文件为空"):
        extract_text(b"", "empty.txt")


def test_docx_roundtrip():
    import docx as docx_lib
    document = docx_lib.Document()
    document.add_paragraph("付款条款：验收合格后 30 日内支付。")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "金额"
    table.rows[0].cells[1].text = "100 万"
    buf = io.BytesIO()
    document.save(buf)
    text, ftype = extract_text(buf.getvalue(), "contract.docx")
    assert ftype == "docx"
    assert "付款条款" in text and "100 万" in text


def test_blank_pdf_reports_no_text():
    from pypdf import PdfReader, PdfWriter
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    _reader = PdfReader(io.BytesIO(buf.getvalue()))  # 确认是合法 PDF
    with pytest.raises(ParserError, match="未能"):
        extract_text(buf.getvalue(), "scan.pdf")
