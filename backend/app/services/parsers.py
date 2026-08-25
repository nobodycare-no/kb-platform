"""文档文本解析器（SRD FR-B01）：txt / md / pdf / docx → 纯文本。

失败必须抛 ParserError 并携带原因，不允许静默返回空结果。
"""
from __future__ import annotations

import io

MAX_SIZE_BYTES = 20 * 1024 * 1024
ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


class ParserError(Exception):
    pass


def detect_extension(filename: str) -> str:
    dot = filename.rfind(".")
    ext = filename[dot:].lower() if dot >= 0 else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise ParserError(f"不支持的文件类型: {ext or '(无扩展名)'}，允许: {sorted(ALLOWED_EXTENSIONS)}")
    return ext


def extract_text(data: bytes, filename: str) -> tuple[str, str]:
    """返回 (text, file_type)。解析失败/空内容抛 ParserError。"""
    if len(data) == 0:
        raise ParserError("文件为空")
    if len(data) > MAX_SIZE_BYTES:
        raise ParserError("文件超过 20MB 上限")

    ext = detect_extension(filename)
    if ext in (".txt", ".md"):
        text = _decode_text(data)
    elif ext == ".pdf":
        text = _extract_pdf(data)
    else:
        text = _extract_docx(data)

    text = text.strip()
    if not text:
        raise ParserError("未能从文件中提取到任何文本（可能是扫描件或空文档）")
    return text, ext[1:]


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ParserError("文本编码无法识别（尝试过 UTF-8 / GB18030）")


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:  # pragma: no cover
        raise ParserError("服务端缺少 pypdf") from e
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
    except Exception as e:
        raise ParserError(f"PDF 解析失败: {e}") from e
    return "\n".join(p for p in pages if p and p.strip())


def _extract_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:  # pragma: no cover
        raise ParserError("服务端缺少 python-docx") from e
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as e:
        raise ParserError(f"DOCX 解析失败: {e}") from e
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
